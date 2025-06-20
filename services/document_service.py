"""
Export Service for generating documents and data exports.

This service handles the creation of:
- PDF documents in FAO analytical brief format
- Word documents with editable content
- Excel files with figure data sheets (fig1_data, fig2_data, etc.)
"""

import logging
import os
import pandas as pd
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

# Document generation imports
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportService:
    """Service for exporting analytical briefs to various formats."""
    
    def __init__(self, config):
        """
        Initialize export service.
        
        Args:
            config: Configuration manager instance
        """
        self.config = config
        self.output_dir = Path(config.get_output_directory()) / "exports"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Set up document styles for PDF
        if REPORTLAB_AVAILABLE:
            self.styles = getSampleStyleSheet()
            self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Set up custom styles for PDF generation."""
        
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=20,
            spaceAfter=20,
            textColor=colors.HexColor('#1f4e79'),
            alignment=TA_CENTER
        ))
        
        # Heading style
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor('#1f4e79'),
            alignment=TA_LEFT
        ))
        
        # Highlight style
        self.styles.add(ParagraphStyle(
            name='Highlight',
            parent=self.styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceBefore=6,
            spaceAfter=6,
            textColor=colors.HexColor('#2c5aa0')
        ))
        
        # Caption style
        self.styles.add(ParagraphStyle(
            name='Caption',
            parent=self.styles['Normal'],
            fontSize=9,
            spaceBefore=6,
            spaceAfter=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#666666')
        ))
        
        # Source style
        self.styles.add(ParagraphStyle(
            name='Source',
            parent=self.styles['Normal'],
            fontSize=8,
            spaceAfter=12,
            textColor=colors.HexColor('#888888')
        ))
    
    def export_pdf(self, brief_results: Dict[str, Any]) -> str:
        """
        Export brief as PDF document.
        
        Args:
            brief_results: Generated brief results
            
        Returns:
            Path to generated PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
        
        # Generate filename
        dataset_name = brief_results['dataset_name'].replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"analytical_brief_{dataset_name}_{timestamp}.pdf"
        filepath = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Build document content
        story = []
        
        # Title page
        story.extend(self._build_title_section(brief_results))
        
        # Highlights
        story.extend(self._build_highlights_section(brief_results))
        
        # Global trends
        story.extend(self._build_global_trends_section(brief_results))
        
        # Regional analysis
        story.extend(self._build_regional_analysis_section(brief_results))
        
        # Country insights
        story.extend(self._build_country_insights_section(brief_results))
        
        # Figures
        story.extend(self._build_figures_section(brief_results))
        
        # Explanatory notes
        story.extend(self._build_explanatory_notes_section(brief_results))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF exported to {filepath}")
        return str(filepath)
    
    def _build_title_section(self, brief_results: Dict[str, Any]) -> List:
        """Build title section for PDF."""
        story = []
        
        # Main title
        title = f"FAOSTAT ANALYTICAL BRIEF\n{brief_results['dataset_name']}"
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 0.3*inch))
        
        # Period
        config = brief_results.get('config', {})
        year_range = config.get('year_range')
        if year_range:
            period = f"{year_range[0]}–{year_range[1]}"
            story.append(Paragraph(period, self.styles['CustomHeading']))
        
        story.append(Spacer(1, 0.5*inch))
        
        return story
    
    def _build_highlights_section(self, brief_results: Dict[str, Any]) -> List:
        """Build highlights section for PDF."""
        story = []
        
        insights = brief_results.get('insights', {})
        highlights = insights.get('highlights', '')
        
        if highlights:
            story.append(Paragraph("HIGHLIGHTS", self.styles['CustomHeading']))
            
            # Convert bullet points to proper format
            highlight_lines = highlights.split('\n')
            for line in highlight_lines:
                if line.strip():
                    # Remove existing bullet markers and add consistent ones
                    clean_line = line.strip().lstrip('→•-').strip()
                    if clean_line:
                        story.append(Paragraph(f"→ {clean_line}", self.styles['Highlight']))
            
            story.append(Spacer(1, 0.3*inch))
        
        return story
    
    def _build_global_trends_section(self, brief_results: Dict[str, Any]) -> List:
        """Build global trends section for PDF."""
        story = []
        
        insights = brief_results.get('insights', {})
        global_trends = insights.get('global_trends', '')
        
        if global_trends:
            story.append(Paragraph("GLOBAL TRENDS", self.styles['CustomHeading']))
            story.append(Paragraph(global_trends, self.styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _build_regional_analysis_section(self, brief_results: Dict[str, Any]) -> List:
        """Build regional analysis section for PDF."""
        story = []
        
        insights = brief_results.get('insights', {})
        regional_analysis = insights.get('regional_analysis', '')
        
        if regional_analysis:
            story.append(Paragraph("REGIONAL ANALYSIS", self.styles['CustomHeading']))
            story.append(Paragraph(regional_analysis, self.styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _build_country_insights_section(self, brief_results: Dict[str, Any]) -> List:
        """Build country insights section for PDF."""
        story = []
        
        insights = brief_results.get('insights', {})
        country_insights = insights.get('country_insights', '')
        
        if country_insights:
            story.append(Paragraph("COUNTRY ANALYSIS", self.styles['CustomHeading']))
            story.append(Paragraph(country_insights, self.styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _build_figures_section(self, brief_results: Dict[str, Any]) -> List:
        """Build figures section for PDF."""
        story = []
        
        figures = brief_results.get('figures', [])
        
        if figures:
            story.append(Paragraph("FIGURES", self.styles['CustomHeading']))
            
            for i, figure in enumerate(figures, 1):
                if os.path.exists(figure['filepath']):
                    try:
                        # Add figure
                        img = Image(figure['filepath'], width=6*inch, height=4*inch)
                        story.append(img)
                        
                        # Add caption
                        caption = f"Figure {i}. {figure['title']}"
                        story.append(Paragraph(caption, self.styles['Caption']))
                        
                        # Add description
                        if figure.get('description'):
                            story.append(Paragraph(figure['description'], self.styles['Normal']))
                        
                        # Add source
                        source_text = "Source: FAO. FAOSTAT. Rome. " + \
                                    datetime.now().strftime('%Y') + \
                                    ". http://www.fao.org/faostat"
                        story.append(Paragraph(source_text, self.styles['Source']))
                        
                        story.append(Spacer(1, 0.3*inch))
                        
                    except Exception as e:
                        logger.warning(f"Could not include figure {figure['filepath']}: {str(e)}")
                        # Add placeholder
                        story.append(Paragraph(f"[Figure {i}: {figure['title']} - Image not available]", self.styles['Caption']))
                        story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _build_explanatory_notes_section(self, brief_results: Dict[str, Any]) -> List:
        """Build explanatory notes section for PDF."""
        story = []
        
        insights = brief_results.get('insights', {})
        explanatory_notes = insights.get('explanatory_notes', '')
        
        if explanatory_notes:
            story.append(Paragraph("EXPLANATORY NOTES", self.styles['CustomHeading']))
            story.append(Paragraph(explanatory_notes, self.styles['Normal']))
        
        return story
    
    def export_word(self, brief_results: Dict[str, Any]) -> str:
        """
        Export brief as Word document.
        
        Args:
            brief_results: Generated brief results
            
        Returns:
            Path to generated Word file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        # Generate filename
        dataset_name = brief_results['dataset_name'].replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"analytical_brief_{dataset_name}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(f"FAOSTAT ANALYTICAL BRIEF\n{brief_results['dataset_name']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Period
        config = brief_results.get('config', {})
        year_range = config.get('year_range')
        if year_range:
            period = doc.add_heading(f"{year_range[0]}–{year_range[1]}", level=1)
            period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Content sections
        insights = brief_results.get('insights', {})
        
        # Highlights
        if insights.get('highlights'):
            doc.add_heading('HIGHLIGHTS', level=1)
            highlight_lines = insights['highlights'].split('\n')
            for line in highlight_lines:
                if line.strip():
                    clean_line = line.strip().lstrip('→•-').strip()
                    if clean_line:
                        p = doc.add_paragraph()
                        p.add_run(f"→ {clean_line}")
        
        # Global trends
        if insights.get('global_trends'):
            doc.add_heading('GLOBAL TRENDS', level=1)
            doc.add_paragraph(insights['global_trends'])
        
        # Regional analysis
        if insights.get('regional_analysis'):
            doc.add_heading('REGIONAL ANALYSIS', level=1)
            doc.add_paragraph(insights['regional_analysis'])
        
        # Country insights
        if insights.get('country_insights'):
            doc.add_heading('COUNTRY ANALYSIS', level=1)
            doc.add_paragraph(insights['country_insights'])
        
        # Figures
        figures = brief_results.get('figures', [])
        if figures:
            doc.add_heading('FIGURES', level=1)
            
            for i, figure in enumerate(figures, 1):
                if os.path.exists(figure['filepath']):
                    try:
                        doc.add_picture(figure['filepath'], width=Inches(6))
                        caption = doc.add_paragraph(f"Figure {i}. {figure['title']}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        if figure.get('description'):
                            doc.add_paragraph(figure['description'])
                        
                        source_text = "Source: FAO. FAOSTAT. Rome. " + \
                                    datetime.now().strftime('%Y') + \
                                    ". http://www.fao.org/faostat"
                        source_p = doc.add_paragraph(source_text)
                        source_p.runs[0].font.size = docx.shared.Pt(8)
                        
                    except Exception as e:
                        logger.warning(f"Could not include figure {figure['filepath']}: {str(e)}")
                        doc.add_paragraph(f"[Figure {i}: {figure['title']} - Image not available]")
        
        # Explanatory notes
        if insights.get('explanatory_notes'):
            doc.add_heading('EXPLANATORY NOTES', level=1)
            doc.add_paragraph(insights['explanatory_notes'])
        
        # Save document
        doc.save(str(filepath))
        
        logger.info(f"Word document exported to {filepath}")
        return str(filepath)
    
    def export_excel_data(self, brief_results: Dict[str, Any]) -> str:
        """
        Export figure data as Excel file with separate sheets.
        
        Args:
            brief_results: Generated brief results
            
        Returns:
            Path to generated Excel file
        """
        # Generate filename
        dataset_name = brief_results['dataset_name'].replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"brief_data_{dataset_name}_{timestamp}.xlsx"
        filepath = self.output_dir / filename
        
        # Create Excel writer
        with pd.ExcelWriter(str(filepath), engine='openpyxl') as writer:
            
            # Summary sheet
            summary_data = self._prepare_summary_sheet(brief_results)
            summary_data.to_excel(writer, sheet_name='Summary', index=False)
            
            # Figure data sheets
            figures = brief_results.get('figures', [])
            config = brief_results.get('config', {})
            
            for i, figure in enumerate(figures, 1):
                sheet_name = f"fig{i}_data"
                figure_data = self._prepare_figure_data(figure, brief_results, config)
                
                if not figure_data.empty:
                    figure_data.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Raw filtered data sheet
            filtered_data = brief_results.get('filtered_data')
            if filtered_data is not None and not filtered_data.empty:
                # Limit to reasonable size for Excel
                if len(filtered_data) > 50000:
                    sample_data = filtered_data.sample(n=50000)
                else:
                    sample_data = filtered_data
                
                sample_data.to_excel(writer, sheet_name='Raw_Data', index=False)
        
        logger.info(f"Excel data exported to {filepath}")
        return str(filepath)
    
    def _prepare_summary_sheet(self, brief_results: Dict[str, Any]) -> pd.DataFrame:
        """Prepare summary information sheet."""
        
        config = brief_results.get('config', {})
        summary_stats = brief_results.get('summary_stats', {})
        
        summary_info = [
            ['Dataset', brief_results.get('dataset_name', '')],
            ['Dataset Code', brief_results.get('dataset_code', '')],
            ['Generated At', brief_results.get('generated_at', '')],
            ['Time Period', f"{config.get('year_range', ['N/A', 'N/A'])[0]} - {config.get('year_range', ['N/A', 'N/A'])[1]}"],
            ['Geographic Scope', config.get('geo_scope', '')],
            ['Selected Items', ', '.join(config.get('selected_items', []))],
            ['Selected Elements', ', '.join(config.get('selected_elements', []))],
            ['Selected Areas', ', '.join(config.get('selected_areas', [])[:10])],  # Limit display
            ['Total Records', summary_stats.get('total_records', 0)],
            ['Areas Covered', summary_stats.get('num_areas', 0)],
            ['Items Covered', summary_stats.get('num_items', 0)],
            ['Elements Covered', summary_stats.get('num_elements', 0)]
        ]
        
        return pd.DataFrame(summary_info, columns=['Attribute', 'Value'])
    
    def _prepare_figure_data(self, figure: Dict[str, Any], brief_results: Dict[str, Any], config: Dict[str, Any]) -> pd.DataFrame:
        """Prepare data for a specific figure."""
        
        df = brief_results.get('filtered_data')
        
        if df is None or df.empty:
            return pd.DataFrame()
        
        figure_type = figure.get('type', '')
        
        try:
            if figure_type == 'trends':
                return self._prepare_trends_figure_data(df, config)
            elif figure_type == 'rankings':
                return self._prepare_rankings_figure_data(df, config)
            elif figure_type == 'composition':
                return self._prepare_composition_figure_data(df, config)
            else:
                # Return a sample of the filtered data
                return df.head(1000)  # Limit size
                
        except Exception as e:
            logger.error(f"Error preparing figure data: {str(e)}")
            return pd.DataFrame()
    
    def _prepare_trends_figure_data(self, df: pd.DataFrame, config: Dict[str, Any]) -> pd.DataFrame:
        """Prepare data for trends figure."""
        
        if 'Year' not in df.columns or 'Value' not in df.columns:
            return df.head(1000)
        
        # Group by year and sum values
        if 'Item' in df.columns and len(df['Item'].unique()) > 1:
            trends_data = df.groupby(['Year', 'Item'])['Value'].sum().reset_index()
        elif 'Element' in df.columns and len(df['Element'].unique()) > 1:
            trends_data = df.groupby(['Year', 'Element'])['Value'].sum().reset_index()
        else:
            trends_data = df.groupby('Year')['Value'].sum().reset_index()
        
        return trends_data
    
    def _prepare_rankings_figure_data(self, df: pd.DataFrame, config: Dict[str, Any]) -> pd.DataFrame:
        """Prepare data for rankings figure."""
        
        if 'Area' not in df.columns or 'Value' not in df.columns:
            return df.head(1000)
        
        # Get latest year data for rankings
        if 'Year' in df.columns:
            latest_year = df['Year'].max()
            rankings_data = df[df['Year'] == latest_year]
        else:
            rankings_data = df
        
        # Group by area and sum values
        area_totals = rankings_data.groupby('Area')['Value'].sum().sort_values(ascending=False)
        
        # Take top 20 for data export
        top_areas = area_totals.head(20).reset_index()
        top_areas['Rank'] = range(1, len(top_areas) + 1)
        
        return top_areas
    
    def _prepare_composition_figure_data(self, df: pd.DataFrame, config: Dict[str, Any]) -> pd.DataFrame:
        """Prepare data for composition figure."""
        
        if 'Value' not in df.columns:
            return df.head(1000)
        
        # Use items if available, otherwise elements
        if 'Item' in df.columns and len(df['Item'].unique()) > 1:
            composition_data = df.groupby('Item')['Value'].sum().sort_values(ascending=False)
            composition_df = composition_data.reset_index()
            composition_df['Percentage'] = (composition_df['Value'] / composition_df['Value'].sum() * 100).round(2)
        elif 'Element' in df.columns and len(df['Element'].unique()) > 1:
            composition_data = df.groupby('Element')['Value'].sum().sort_values(ascending=False)
            composition_df = composition_data.reset_index()
            composition_df['Percentage'] = (composition_df['Value'] / composition_df['Value'].sum() * 100).round(2)
        else:
            return df.head(1000)
        
        return composition_df